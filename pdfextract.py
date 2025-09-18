import fitz
import pdfplumber
from docx import Document
from docx.shared import Inches
import google.generativeai as genai
import os

genai.configure(api_key="API key")

def gemini_ai(sometext, imgpath=None):
    m = genai.GenerativeModel("gemini-2.0-flash")
    try:
        if imgpath:
            res = m.generate_content([sometext, genai.upload_file(imgpath)])
        else:
            res = m.generate_content(sometext)
        if res and res.text:
            return res.text
        else:
            return "no ai answer"
    except Exception as er:
        return f"error {er}"

def cleanline(txt):
    return txt.replace("*", "").replace("#", "").replace("-", "").replace("•", "").strip()

def put_ai_text(docfile, txt):
    lines = txt.split("\n")
    now = None
    for l in lines:
        l = cleanline(l)
        if not l:
            continue
        low = l.lower()
        if low.startswith("context"):
            docfile.add_heading("Context", level=3)
            now = "context"
        elif low.startswith("key points"):
            docfile.add_heading("Key Points", level=3)
            now = "points"
        elif low.startswith("implications"):
            docfile.add_heading("Implications", level=3)
            now = "imp"
        else:
            if now == "points":
                docfile.add_paragraph(l, style="List Bullet")
            else:
                docfile.add_paragraph(l)

def makeReport(pdfname, outname="ai_report.docx"):
    d = Document()
    p = fitz.open(pdfname)
    made_imgs = []
    with pdfplumber.open(pdfname) as pp:
        for pageNo in range(len(p)):
            pg = p[pageNo]
            plpg = pp.pages[pageNo]
            theText = pg.get_text("text")

            d.add_heading("Page " + str(pageNo+1), level=1)

            if theText.strip():
                d.add_heading("Original Text:", level=2)
                d.add_paragraph(theText)
                ans = gemini_ai(
                    f"""Analyze this text from page {pageNo+1}.

                    Context: say what it is about.
                    Key Points: main details and insights.
                    Implications: meaning or importance.

                    Text:
                    {theText}"""
                )
                d.add_heading("Detailed Explanation (AI):", level=2)
                put_ai_text(d, ans)

            imgs = pg.get_images(full=True)
            for i, im in enumerate(imgs):
                xref = im[0]
                base = p.extract_image(xref)
                data = base["image"]
                ext = base["ext"]
                fname = f"pg{pageNo+1}_img{i}.{ext}"
                with open(fname, "wb") as f:
                    f.write(data)
                made_imgs.append(fname)
                d.add_heading("Image " + str(i+1) + " from Page " + str(pageNo+1), level=2)
                d.add_picture(fname, width=Inches(5))
                imans = gemini_ai(
                    f"""Analyze this image/graph from page {pageNo+1}.

                    Context: what is shown.
                    Key Points: important things or patterns.
                    Implications: why it matters.
                    """,
                    imgpath=fname
                )
                d.add_heading("Image Analysis (AI):", level=2)
                put_ai_text(d, imans)

            tbls = plpg.extract_tables()
            for tnum, t in enumerate(tbls):
                d.add_heading("Table " + str(tnum+1) + " from Page " + str(pageNo+1), level=2)
                tb = d.add_table(rows=len(t), cols=len(t[0]))
                for r, row in enumerate(t):
                    for c, cell in enumerate(row):
                        tb.cell(r, c).text = str(cell) if cell else ""
                tabtxt = "\n".join([", ".join(r) for r in t])
                tabsum = gemini_ai(
                    f"""Analyze this table from page {pageNo+1}.

                    Context: what the table is about.
                    Key Points: main values and findings.
                    Implications: meaning or impact.

                    Table:
                    {tabtxt}"""
                )
                d.add_heading("Table Analysis (AI):", level=2)
                put_ai_text(d, tabsum)
    d.save(outname)
    for f in made_imgs:
        try:
            os.remove(f)
        except:
            pass
    print("AI report saved to", outname)

if __name__ == "__main__":
    makeReport("input.pdf", "ai_output.docx")

